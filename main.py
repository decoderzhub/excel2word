import streamlit as st
import openpyxl
# from docxtpl import DocxTemplate, RichText
from docx import Document


def populate_word_template(excel_file, word_template, output_file, excel_data, question_header, answer_header):
    workbook = openpyxl.load_workbook(excel_file)
    sheet = workbook[excel_data['sheet_name']]

    # Find the column indices based on the header names
    question_index, answer_index = None, None
    # Assuming headers are in the first row
    for col_idx, cell in enumerate(sheet[1], start=1):
        if cell.value == question_header:
            question_index = col_idx
        if cell.value == answer_header:
            answer_index = col_idx

    if question_index is None or answer_index is None:
        raise ValueError(
            f"Column headers '{question_header}' or '{answer_header}' not found.")

    # Retrieve all rows with values in the specified columns
    questions, answers = [], []
    for row in sheet.iter_rows(min_row=2):
        question_cell_value = row[question_index - 1].value
        answer_cell_value = row[answer_index - 1].value

        if question_cell_value and answer_cell_value:
            questions.append(
                str(question_cell_value).strip().replace("\n", ""))
            answers.append(str(answer_cell_value).strip().replace('\n', ""))

     # Load the Word template and find the content control
    document = Document(word_template)

    # Insert question-answer pairs and apply styles
    for question, answer in zip(questions, answers):
        # Insert the question and apply the style
        document.add_paragraph(question, style="Heading 1")

        # Insert the answer and apply the style
        document.add_paragraph(answer, style="Heading 2")

    # Save the output file
    document.save(output_file)


# Streamlit app
st.title("Excel to Word Populator")

uploaded_excel_file = st.file_uploader(
    "Upload Excel file", type=["xlsx", "xls"])

if uploaded_excel_file is not None:
    excel_file = "uploaded_excel.xlsx"
    with open(excel_file, "wb") as f:
        f.write(uploaded_excel_file.getbuffer())

    st.success("Excel file uploaded")

    # Let the user choose the sheet
    st.subheader("Select sheet")
    sheet_names = openpyxl.load_workbook(excel_file).sheetnames
    selected_sheet_name = st.selectbox("Select sheet", sheet_names)

    # Load the selected sheet
    workbook = openpyxl.load_workbook(excel_file)
    sheet = workbook[selected_sheet_name]

    # Extract headers from the first row
    headers = [cell.value for cell in sheet[1]]

    # Let the user choose the headers to iterate through
    st.subheader("Select column headers")
    selected_question_header = st.selectbox(
        "Select question header", headers, key="question_header")
    selected_answer_header = st.selectbox(
        "Select answer header", headers, key="answer_header")

if st.button("Generate Word document"):
    # Replace with the path to your Word template and output file
    word_template = "word_template.docx"
    output_file = "output_word_file.docx"

    excel_data = {
        "sheet_name": selected_sheet_name,
    }

    populate_word_template(excel_file, word_template, output_file,
                           excel_data, selected_question_header, selected_answer_header)
    st.success("Word document generated")

    # Display the generated Word document for download
    with open(output_file, "rb") as f:
        output_file_bytes = f.read()
    st.download_button("Download Word document", output_file_bytes, file_name="output_word_file.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
